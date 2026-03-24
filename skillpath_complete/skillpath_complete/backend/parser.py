"""
Document Parser Module
Handles PDF, DOCX, and TXT extraction for resume ingestion.
"""

import io
import re


def extract_text_from_file(file_storage):
    """
    Extract plain text from an uploaded file (PDF, DOCX, or TXT).
    Returns cleaned text string.
    """
    filename = file_storage.filename.lower()
    file_bytes = file_storage.read()

    if filename.endswith(".pdf"):
        return _extract_from_pdf(file_bytes)
    elif filename.endswith(".docx"):
        return _extract_from_docx(file_bytes)
    elif filename.endswith(".txt"):
        return _extract_from_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: '{filename}'. Please upload PDF, DOCX, or TXT.")


def _extract_from_pdf(file_bytes: bytes) -> str:
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(io.BytesIO(file_bytes))
        return _clean_text(text)
    except Exception as e:
        raise RuntimeError(f"PDF extraction failed: {str(e)}. Try converting to DOCX or TXT.")


def _extract_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return _clean_text("\n".join(paragraphs))
    except Exception as e:
        raise RuntimeError(f"DOCX extraction failed: {str(e)}")


def _extract_from_txt(file_bytes: bytes) -> str:
    try:
        text = file_bytes.decode("utf-8", errors="replace")
        return _clean_text(text)
    except Exception as e:
        raise RuntimeError(f"TXT extraction failed: {str(e)}")


def _clean_text(text: str) -> str:
    """Remove excessive whitespace and non-printable characters."""
    if not text:
        return ""
    # Normalize whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove non-printable characters but keep common punctuation
    text = re.sub(r'[^\x20-\x7E\n]', ' ', text)
    # Collapse multiple newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()
